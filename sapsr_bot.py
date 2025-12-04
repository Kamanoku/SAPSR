import asyncio
import os
import logging
import re
from datetime import datetime
from typing import List, Dict, Tuple, Any

# Библиотеки для Telegram
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup

# Библиотеки для работы с документами
import docx
import PyPDF2

# ============================================================
#  AGENT 1: PERCEPTION AGENT (УЛУЧШЕННЫЙ PDF ПАРСИНГ)
# ============================================================
class PerceptionAgent:
    @staticmethod
    def _normalize_text(s: str) -> str:
        if s is None: return ""
        # 1. Замена спецсимволов и подчеркиваний
        s = s.replace("_", " ").replace("\u00A0", " ").replace("\u200B", "").replace("\uFEFF", "")
        # 2. Замена переносов и табуляций
        s = s.replace('\n', ' ').replace('\r', ' ').replace('\t', ' ')
        # 3. FIX: Удаление "разрывов" слов в PDF (сентяб ря -> сентября)
        # Если видим "буква-пробел-буква", и это не предлог (эвристика)
        # (Это базовая защита, идеальная требует словаря)
        # Здесь мы просто схлопываем множественные пробелы
        s = re.sub(r"[ \t\v\f]+", " ", s)
        return s.strip()

    @staticmethod
    def load_content(path: str) -> List[str]:
        lower = path.lower()
        if lower.endswith(".docx"):
            return PerceptionAgent._load_docx(path)
        elif lower.endswith(".pdf"):
            return PerceptionAgent._load_pdf(path)
        else:
            raise ValueError("Формат не поддерживается")

    @staticmethod
    def _load_docx(path: str) -> List[str]:
        doc = docx.Document(path)
        paragraphs = []
        for p in doc.paragraphs:
            paragraphs.append(PerceptionAgent._normalize_text(p.text))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(PerceptionAgent._normalize_text(cell.text))
        return paragraphs

    @staticmethod
    def _load_pdf(path: str) -> List[str]:
        lines = []
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text = page.extract_text()
                if not text: continue
                
                # FIX: PyPDF2 часто разбивает таблицы на короткие строки.
                # Мы не нормализуем сразу, а отдаем сырые строки,
                # но SchemaAgent теперь умеет их склеивать.
                for ln in text.splitlines():
                    norm_ln = PerceptionAgent._normalize_text(ln)
                    if norm_ln:
                        lines.append(norm_ln)
        return lines

# ============================================================
#  AGENT 2: SCHEMA AGENT (FIX: СКЛЕЙКА РАЗОРВАННЫХ ТЕГОВ)
# ============================================================
class SchemaAgent:
    def parse_template(self, paragraphs: List[str]) -> List[Dict]:
        placeholders = []
        # Regex теперь допускает пробелы внутри тега на случай разрывов
        pattern = re.compile(
            r"\[\[\s*([^:\]\n]+?)\s*:\s*([^,:]\s*[^,\]\n]+?)"
            r"(?:\s*:\s*([^:\]\n]+?)\s*:\s*([^,\]\n]+?))?"
            r"(?:\s*,\s*(optional))?\s*\]\]",
            flags=re.IGNORECASE,
        )

        seen_names = set()
        
        # БУФЕР ДЛЯ СКЛЕЙКИ СТРОК
        # Если строка содержит '[[', но не содержит ']]', мы не обрабатываем её,
        # а ждем следующую, чтобы склеить. Это чинит таблицы в PDF.
        buffer = ""
        buffer_start_idx = 0

        processed_paragraphs = [] # Список (текст, оригинальный индекс)

        # 1. Предварительная обработка: склейка разорванных тегов
        for idx, para in enumerate(paragraphs):
            # Если в буфере что-то есть, добавляем текущую строку к нему
            if buffer:
                buffer += " " + para
                # Если тег закрылся
                if "]]" in para:
                    processed_paragraphs.append((buffer, buffer_start_idx))
                    buffer = ""
                continue
            
            # Если начало тега есть, а конца нет — начинаем накапливать буфер
            if "[[" in para and "]]" not in para:
                buffer = para
                buffer_start_idx = idx
                continue
            
            # Обычная строка
            processed_paragraphs.append((para, idx))

        # 2. Парсинг
        for text, original_idx in processed_paragraphs:
            if not text.strip(): continue
            
            for m in pattern.finditer(text):
                name = m.group(1).strip()
                if name.lower() in seen_names: continue
                seen_names.add(name.lower())

                data = {
                    "name": name,
                    "type": self._normalize_type(m.group(2).strip()),
                    "group_name": m.group(3).strip() if m.group(3) else "",
                    "group_condition": m.group(4).strip() if m.group(4) else "",
                    "optional": bool(m.group(5)),
                    "anchor_before": "",
                }

                # Якорь берем из склеенного текста
                left_part = text[: m.start()].strip()
                if left_part:
                    # Берем последние 30 символов, чтобы якорь не был слишком длинным
                    data["anchor_before"] = left_part[-40:]
                else:
                    # Ищем в предыдущих "чистых" параграфах
                    # (Упрощенная логика для склеенных строк)
                    stop_patterns = ["утверждаю", "задание", "введение"]
                    for j in range(len(processed_paragraphs) - 1, -1, -1):
                        prev_txt, prev_idx = processed_paragraphs[j]
                        if prev_idx >= original_idx: continue # Не смотрим вперед
                        
                        if prev_txt and "[[" not in prev_txt and not any(s in prev_txt.lower() for s in stop_patterns):
                            data["anchor_before"] = prev_txt
                            break

                placeholders.append(data)

        return placeholders

    def _normalize_type(self, t):
        t = t.lower()
        if t in ("int", "integer", "num", "number"): return "number"
        if t in ("str", "string", "text"): return "string"
        return t

# ============================================================
#  AGENT 3: EXTRACTION AGENT (КРИТИЧЕСКИЕ ИСПРАВЛЕНИЯ)
# ============================================================
class ExtractionAgent:
    def __init__(self):
        self.stop_words = [
            "введение", "заключение", "список использованных источников",
            "приложение", "задание", "руководитель", "куратор", "проверяющий",
            "наименование этапов" # Добавлено, чтобы не хватало заголовки таблиц
        ]

    def _clean_anchor(self, text):
        return re.sub(r'[^a-zA-Zа-яА-Я0-9]', '', text).lower()

    def find_value(self, item: Dict, doc_paragraphs: List[str], start_cursor: int) -> Tuple[bool, str, int]:
        anchor = item.get("anchor_before")
        expected_type = item.get("type")
        
        if not anchor:
            return False, None, start_cursor

        clean_anchor = self._clean_anchor(anchor)
        
        # Поиск строки, содержащей якорь
        target_idx = -1
        for i in range(start_cursor, len(doc_paragraphs)):
            # Нестрогое сравнение: содержится ли очищенный якорь в очищенной строке
            if clean_anchor in self._clean_anchor(doc_paragraphs[i]):
                target_idx = i
                break
        
        if target_idx == -1:
            return False, None, start_cursor

        # === СТРАТЕГИЯ 0: ИЗВЛЕЧЕНИЕ ИЗ ТОЙ ЖЕ СТРОКИ (Same Line) ===
        # Это решит проблему "3. Разработка фыв" и "Сальников"
        current_text = doc_paragraphs[target_idx]
        
        # Находим, где кончается якорь в реальном тексте
        # (Упрощенно: разбиваем строку по тексту якоря, если он там есть целиком)
        # Для надежности используем regex escape
        match = re.search(re.escape(PerceptionAgent._normalize_text(anchor)), current_text, re.IGNORECASE)
        
        candidate_same_line = ""
        if match:
            candidate_same_line = current_text[match.end():].strip()
        else:
            # Если точного совпадения нет (из-за очистки), пробуем эвристику:
            # Берем последние 70% строки, если она длинная
            pass 

        if candidate_same_line:
            # Если ожидаем число
            if expected_type == "number":
                num = self._extract_number(candidate_same_line)
                if num: return True, num, target_idx
            # Если ожидаем строку/дату
            else:
                # Проверка: не является ли остаток просто мусором или стоп-словом
                if len(candidate_same_line) > 1 and not any(sw in candidate_same_line.lower() for sw in self.stop_words):
                    return True, candidate_same_line, target_idx

        # === СТРАТЕГИЯ 1: ПОИСК В СЛЕДУЮЩИХ СТРОКАХ ===
        # Если на текущей строке пусто, смотрим вниз (макс 4 строки)
        for k in range(target_idx + 1, min(len(doc_paragraphs), target_idx + 5)):
            cand = doc_paragraphs[k].strip()
            if not cand: continue
            
            # Стоп-факторы
            cand_lower = cand.lower()
            if any(sw in cand_lower for sw in self.stop_words): break
            if re.match(r"^\d+\.", cand): break # Следующий пункт списка (например, 4. Заключение)

            if expected_type == "number":
                num = self._extract_number(cand)
                if num: return True, num, k
            else:
                return True, cand, k

        return False, None, target_idx

    def _extract_number(self, text):
        # Ищет число, игнорируя окружающий текст
        m = re.search(r"([+-]?\s*\d+([.,]\d+)?)", text)
        if m: return m.group(1).replace(" ", "")
        return None

# ============================================================
#  AGENT 4: ANALYST AGENT (Агент-Аналитик)
#  Роль: Валидация типов, Математический анализ, Логические выводы.
# ============================================================

class AnalystAgent:
    def validate_type(self, value: str, expected_type: str) -> bool:
        if not value: return False
        v = value.strip()
        if expected_type == "string": 
            return bool(re.search(r"[A-Za-zА-Яа-яЁё]", v))
        if expected_type == "number": 
            # Строгая проверка числа
            clean_v = v.replace(' ', '').replace(',', '.')
            try:
                float(clean_v)
                return True
            except ValueError:
                return False
        if expected_type == "date":
            return bool(re.search(r"\d{1,2}[\.\s][\w\.]+\s?\d{4}", v))
        return True

    def analyze_groups(self, extraction_results: List[Dict]) -> List[Dict]:
        """
        Интеллектуальный анализ: группировка данных, вычисление сумм/средних
        и проверка условий (например, SUM=100).
        """
        groups = {}
        # Сбор данных
        for res in extraction_results:
            g_name = res.get("group_name")
            g_cond = res.get("group_condition")
            if not g_name or not g_cond: continue
            
            key = (g_name, g_cond)
            if key not in groups:
                groups[key] = {"values": [], "missing": []}
            
            if res["status"] == "ok" and res["expected_type"] == "number":
                try:
                    val = float(res["value"].replace(',', '.').replace(' ', ''))
                    groups[key]["values"].append(val)
                except:
                    groups[key]["missing"].append(res["field"])
            else:
                groups[key]["missing"].append(res["field"])

        # Вычисление и генерация выводов
        analysis_report = []
        for (name, condition), data in groups.items():
            if data["missing"]:
                analysis_report.append({
                    "type": "group_error",
                    "msg": f"⚠️ Группа '{name}': Невозможно проверить условие '{condition}'. Ошибки в полях: {', '.join(data['missing'])}"
                })
                continue

            # Парсинг условия: (SUM|AVG)([<=>!]+)(\d+)
            m = re.match(r"(SUM|AVG)([<=>!]+)(\d+(\.\d+)?)", condition.upper().replace(' ', ''))
            if not m:
                analysis_report.append({"type": "group_error", "msg": f"❌ Группа '{name}': Некорректный синтаксис условия '{condition}'"})
                continue
            
            op_type, operator, target_str = m.group(1), m.group(2), m.group(3)
            target = float(target_str)
            
            # Вычисления
            calculated = sum(data["values"])
            if op_type == "AVG" and data["values"]:
                calculated /= len(data["values"])
            
            # Логическая проверка
            valid = False
            if operator == '=': valid = abs(calculated - target) < 0.01
            elif operator == '>': valid = calculated > target
            elif operator == '<': valid = calculated < target
            elif operator == '>=': valid = calculated >= target
            elif operator == '<=': valid = calculated <= target
            
            icon = "✅" if valid else "❌"
            result_text = "соответствует" if valid else "НЕ соответствует"
            
            analysis_report.append({
                "type": "group_result",
                "valid": valid,
                "msg": f"{icon} Группа '{name}': {op_type} = {calculated:.2f}. Это {result_text} условию {condition}."
            })
            
        return analysis_report

# ============================================================
#  SYSTEM: COORDINATOR (Оркестратор)
# ============================================================

class MultiAgentCheckSystem:
    def __init__(self):
        self.perceptor = PerceptionAgent()
        self.schema = SchemaAgent()
        self.extractor = ExtractionAgent()
        self.analyst = AnalystAgent()

    def process(self, template_path: str, doc_path: str) -> str:
        try:
            # 1. Восприятие
            tpl_paras = self.perceptor.load_content(template_path)
            doc_paras = self.perceptor.load_content(doc_path)

            # 2. Построение схемы
            plan = self.schema.parse_template(tpl_paras)
            if not plan:
                return "❌ Ошибка: В шаблоне не найдено тегов вида [[name:type]]."

            # 3. Извлечение и первичная проверка
            results = []
            cursor = 0
            # Пропуск пустых строк в начале документа
            while cursor < len(doc_paras) and not doc_paras[cursor].strip():
                cursor += 1
            
            for item in plan:
                found, val, idx = self.extractor.find_value(item, doc_paras, cursor)
                
                res = {
                    "field": item["name"],
                    "expected_type": item["type"],
                    "group_name": item["group_name"],
                    "group_condition": item["group_condition"],
                    "value": val,
                    "optional": item["optional"]
                }
                
                if found:
                    is_valid = self.analyst.validate_type(val, item["type"])
                    res["status"] = "ok" if is_valid else "type_error"
                    # Сдвигаем курсор, но не слишком агрессивно, если это таблица
                    cursor = max(cursor, idx) 
                else:
                    res["status"] = "missing_optional" if item["optional"] else "missing"
                
                results.append(res)

            # 4. Интеллектуальный анализ (Группы и математика)
            group_analysis = self.analyst.analyze_groups(results)

            # 5. Генерация отчета
            return self._generate_human_report(doc_path, results, group_analysis)

        except Exception as e:
            logging.error(f"System Error: {e}", exc_info=True)
            return f"🔥 Критическая ошибка системы: {str(e)}"

    def _generate_human_report(self, doc_name, results, group_analysis):
        lines = [f"🤖 <b>Результаты проверки</b>", f"📄 Файл: {os.path.basename(doc_name)}", ""]
        
        lines.append("<b>1. Проверка полей:</b>")
        for r in results:
            if r["status"] == "ok":
                lines.append(f"✅ <b>{r['field']}</b>: {r['value']}")
            elif r["status"] == "type_error":
                lines.append(f"⚠️ <b>{r['field']}</b>: '{r['value']}' (Неверный тип, жду {r['expected_type']})")
            elif r["status"] == "missing":
                lines.append(f"❌ <b>{r['field']}</b>: Не найдено")
            elif r["status"] == "missing_optional":
                lines.append(f"ℹ️ {r['field']}: пропущено (необяз.)")

        if group_analysis:
            lines.append("\n<b>2. Логический анализ:</b>")
            for ga in group_analysis:
                lines.append(ga["msg"])
        
        return "\n".join(lines)

# ============================================================
#  TELEGRAM BOT LOGIC
# ============================================================

BOT_TOKEN = "8124707173:AAEUWIG6cU8ErdX_ItQZdbWNGD3JRLwjjNo" # <-- Вставьте токен

logging.basicConfig(level=logging.INFO)
bot = Bot(token=BOT_TOKEN)
dp = Dispatcher(storage=MemoryStorage())
system = MultiAgentCheckSystem()

TEMP_DIR = "temp_files"
os.makedirs(TEMP_DIR, exist_ok=True)

class Workflow(StatesGroup):
    waiting_for_template = State()
    waiting_for_document = State()

@dp.message(Command("start"))
async def cmd_start(message: types.Message, state: FSMContext):
    await message.answer(
        "👋 Привет! Я <b>САПСР</b>.\n\n"
        "Пришлите <b>ШАБЛОН</b> в формате docx/pdf\n",
        parse_mode="HTML"
    )
    await state.set_state(Workflow.waiting_for_template)

@dp.message(Workflow.waiting_for_template, F.document)
async def process_template(message: types.Message, state: FSMContext):
    file_name = message.document.file_name
    if not (file_name.endswith('.docx') or file_name.endswith('.pdf')):
        await message.answer("❌ Поддерживаются только .docx и .pdf файлы.")
        return

    file = await bot.get_file(message.document.file_id)
    local_path = os.path.join(TEMP_DIR, f"tpl_{message.from_user.id}_{file_name}")
    await bot.download_file(file.file_path, local_path)
    
    await state.update_data(template_path=local_path)
    await message.answer(f"✅ Шаблон <b>{file_name}</b> загружен. \nЖду документ для проверки.", parse_mode="HTML")
    await state.set_state(Workflow.waiting_for_document)

@dp.message(Workflow.waiting_for_document, F.document)
async def process_document(message: types.Message, state: FSMContext):
    data = await state.get_data()
    template_path = data.get("template_path")
    if not template_path:
        await message.answer("⚠️ Шаблон потерян. Начните с /start")
        return

    msg = await message.answer("⏳ Агенты обрабатывают данные...")
    
    file_name = message.document.file_name
    file = await bot.get_file(message.document.file_id)
    doc_path = os.path.join(TEMP_DIR, f"doc_{message.from_user.id}_{file_name}")
    await bot.download_file(file.file_path, doc_path)
    
    # Запуск системы в отдельном потоке
    report = await asyncio.to_thread(system.process, template_path, doc_path)
    
    # Разбивка на части, если сообщение слишком длинное для Telegram (4096 символов)
    if len(report) > 4000:
        for x in range(0, len(report), 4000):
            await message.answer(report[x:x+4000], parse_mode="HTML")
    else:
        await msg.edit_text(report, parse_mode="HTML")
    
    await message.answer("Можете прислать следующий документ или /start для смены шаблона.")

@dp.message(Command("cancel"))
async def cmd_cancel(message: types.Message, state: FSMContext):
    await state.clear()
    await message.answer("Сброс выполнен. Жмите /start")

async def main():
    await bot.delete_webhook(drop_pending_updates=True)
    await dp.start_polling(bot)

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        pass
