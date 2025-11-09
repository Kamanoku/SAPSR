# -*- coding: utf-8 -*-
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import docx
import PyPDF2
import os
import re
from datetime import datetime


# ============================================================
#  DocumentLoader
# ============================================================
class DocumentLoader:
    """Загружает текст и параграфы из .docx и .pdf файлов."""

    @staticmethod
    def _normalize_text(s: str) -> str:
        if s is None:
            return ""
        s = s.replace("\u00A0", " ").replace("\u200B", "").replace("\uFEFF", "")
        s = re.sub(r"[ \t\v\f\u00A0]+", " ", s)
        return s.strip()

    # ---------- обновлённые функции ----------
    @staticmethod
    def load_docx_text_and_paragraphs(path: str, dedupe: bool = True, preserve_empty: bool = False):
        """
        Загружает параграфы из .docx.
        - dedupe=True : убирает точные дубликаты (старое поведение)
        - preserve_empty=True : сохраняет пустые параграфы как "" (по умолчанию их удаляем)
        """
        doc = docx.Document(path)
        paragraphs = []
        seen = set()

        def add_para(text):
            if text is None:
                text = ""
            t_norm = DocumentLoader._normalize_text(text) if text else ""
            if t_norm == "" and not preserve_empty:
                return
            if dedupe:
                if t_norm and t_norm not in seen:
                    paragraphs.append(t_norm)
                    seen.add(t_norm)
                elif t_norm == "" and preserve_empty:
                    paragraphs.append(t_norm)
            else:
                if t_norm == "" and not preserve_empty:
                    return
                paragraphs.append(t_norm)

        for p in doc.paragraphs:
            text = "".join(run.text for run in p.runs)
            add_para(text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        text = "".join(run.text for run in p.runs)
                        add_para(text)

        full_text = "\n".join(paragraphs)
        return full_text, paragraphs

    @staticmethod
    def load_pdf_text_and_paragraphs(path: str, dedupe: bool = True, preserve_empty: bool = False):
        """
        Загружает строки/параграфы из PDF.
        PyPDF2.extract_text возвращает текст страницы — разбиваем по строкам.
        """
        text_lines = []
        seen = set()

        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text() or ""
                for ln in page_text.splitlines():
                    ln_norm = DocumentLoader._normalize_text(ln)
                    if ln_norm == "" and not preserve_empty:
                        continue
                    if dedupe:
                        if ln_norm and ln_norm not in seen:
                            text_lines.append(ln_norm)
                            seen.add(ln_norm)
                        elif ln_norm == "" and preserve_empty:
                            text_lines.append(ln_norm)
                    else:
                        if ln_norm == "" and not preserve_empty:
                            continue
                        text_lines.append(ln_norm)

        full_text = "\n".join(text_lines)
        return full_text, text_lines

    @staticmethod
    def get_paragraphs(path: str):
        lower = path.lower()
        if lower.endswith(".docx"):
            # Для проверки документа сохраняем пустые параграфы и не удаляем дубликаты
            _, paras = DocumentLoader.load_docx_text_and_paragraphs(path, dedupe=False, preserve_empty=True)
            return paras
        elif lower.endswith(".pdf"):
            _, paras = DocumentLoader.load_pdf_text_and_paragraphs(path, dedupe=False, preserve_empty=True)
            return paras
        else:
            raise ValueError("Поддерживаются только .docx и .pdf")


# ============================================================
#  Template
# ============================================================
class Template:
    """Хранит список placeholders и их anchors."""

    def __init__(self, placeholders=None, source_path=None):
        self.placeholders = placeholders or []
        self.source_path = source_path

    def get_placeholders(self):
        return self.placeholders

    @staticmethod
    def _normalize_type(raw_type: str) -> str:
        t = raw_type.strip().lower()
        if t in ("int", "integer", "num", "number", "float"):
            return "number"
        if t in ("str", "string", "text"):
            return "string"
        if t in ("date", "dt"):
            return "date"
        return t

    @staticmethod
    def extract_placeholders_from_paragraphs(paragraphs: list) -> list:
        """
        Находит placeholders и вычисляет anchor_before и anchor_after.
        Исправленные алгоритмы:
        - ограничена дистанция между placeholders (max_template_distance)
        - игнорируются параграфы, содержащие [[...]] как anchors
        - при выборе anchor_after пропускаются служебные строки вроде "(подпись)" и т.п.
        """
        placeholders = []
        inline_pattern = re.compile(
            r"\[\[\s*([^:\]\n]+?)\s*:\s*([^,\]\n]+?)(?:\s*,\s*(optional))?\s*\]\]",
            flags=re.IGNORECASE,
        )

        skip_patterns = [
            "утверждаю",
            "задание",
            "введение",
            "заключение",
            "список использованных источников",
            "примерный календарный график",
            "подпись обучающегося",
        ]

        for idx, para in enumerate(paragraphs):
            if not para.strip():
                continue

            for m in inline_pattern.finditer(para):
                raw_name = m.group(1).strip()
                raw_type = m.group(2).strip()
                optional_flag = bool(m.group(3))

                # ---------------- anchor_before ----------------
                left_part = para[: m.start()].strip()
                if left_part:
                    anchor_before = left_part
                else:
                    anchor_before = ""
                    for j in range(idx - 1, -1, -1):
                        prev_para = paragraphs[j].strip()
                        if (
                            prev_para
                            and not inline_pattern.search(prev_para)
                            and not any(sp in prev_para.lower() for sp in skip_patterns)
                        ):
                            anchor_before = prev_para
                            break

                # ---------------- anchor_after ----------------
                right_part = para[m.end() :].strip()
                if right_part:
                    anchor_after = right_part
                else:
                    anchor_after = ""
                    max_template_distance = 6  # увеличиваем дальность поиска anchor_after
                    forbidden_after = [
                        "(подпись)",
                        "(инициалы",
                        "фамилия",
                        "подпись",
                        "инициалы",
                        "(инициалы, фамилия)",
                        "подпись обучающегося",
                        "(подпись обучающегося)",
                    ]
                    for j in range(idx + 1, min(len(paragraphs), idx + 1 + max_template_distance)):
                        next_para = paragraphs[j].strip()
                        if not next_para:
                            continue
                        next_lower = next_para.lower()
                        # пропускаем служебные строки — подписи и т.д.
                        if any(f in next_lower for f in forbidden_after):
                            continue
                        if (
                            not inline_pattern.search(next_para)
                            and not any(sp in next_lower for sp in skip_patterns)
                        ):
                            anchor_after = next_para
                            break

                placeholders.append(
                    {
                        "name": raw_name,
                        "type": Template._normalize_type(raw_type),
                        "optional": optional_flag,
                        "anchor_before": anchor_before,
                        "anchor_after": anchor_after,
                        "source_paragraph": para,
                        "para_index": idx,
                    }
                )

        # удаление дубликатов
        seen = set()
        unique = []
        for p in placeholders:
            key = (
                p["name"].lower(),
                re.sub(r"\s+", " ", p["anchor_before"].strip()).lower()
                if p["anchor_before"]
                else "",
                re.sub(r"\s+", " ", p["anchor_after"].strip()).lower()
                if p["anchor_after"]
                else "",
            )
            if key not in seen:
                seen.add(key)
                unique.append(p)

        # маркер "следующий placeholder подряд"
        for i in range(len(unique) - 1):
            if unique[i + 1]["para_index"] - unique[i]["para_index"] <= 1:
                unique[i]["next_is_placeholder"] = True
            else:
                unique[i]["next_is_placeholder"] = False

        return unique

    @classmethod
    def load_from_file(cls, path: str):
        lower = path.lower()
        if lower.endswith(".docx"):
            _, paragraphs = DocumentLoader.load_docx_text_and_paragraphs(path, dedupe=True, preserve_empty=False)
        elif lower.endswith(".pdf"):
            _, paragraphs = DocumentLoader.load_pdf_text_and_paragraphs(path, dedupe=True, preserve_empty=False)
        else:
            raise ValueError("Поддерживаются только .docx и .pdf")

        placeholders = cls.extract_placeholders_from_paragraphs(paragraphs)
        if not placeholders:
            raise ValueError("В шаблоне не найдено ни одного заполнителя [[...]]")
        return cls(placeholders=placeholders, source_path=path)


# ============================================================
#  DocumentChecker
# ============================================================
class DocumentChecker:
    """Проверяет документ по anchors."""

    def __init__(self, template: Template):
        self.template = template

    @staticmethod
    def _validate_type(value: str, expected_type: str) -> bool:
        if not value:
            return False
        v = value.strip()
        if expected_type == "string":
            return bool(re.search(r"[A-Za-zА-Яа-яЁё]", v))
        if expected_type == "number":
            return bool(re.fullmatch(r"[+-]?\d+([.,]\d+)?", v))
        if expected_type == "date":
            return bool(
                re.fullmatch(r"\d{1,2}\.\d{1,2}\.\d{4}", v)
                or re.search(r"\d{1,2}\s+[А-Яа-яёЁ]+\.?\s+\d{4}", v)
            )
        return True

    @staticmethod
    def _is_anchor_like(value: str, anchors: list) -> bool:
        if not value:
            return False
        v = re.sub(r"\s+", " ", value).strip().lower()
        for a in anchors:
            if not a:
                continue
            a_norm = re.sub(r"\s+", " ", a).strip().lower()
            if v == a_norm:
                return True
        return False

    def _find_value_using_anchors(
        self, anchor_before, anchor_after, doc_paragraphs, start_index=0, expected_type=None, next_is_placeholder=False
    ):
        """Поиск значения между anchors с ограничением дистанции и без глобального fallback."""

        stop_words = [
            "введение",
            "заключение",
            "список использованных источников",
            "примерный календарный график",
            "приложение",
            "руководитель курсового проекта",
            "куратор",
            "проверяющий",
            "обучающемуся",
            "задание",
        ]

        def find_positions(anchor):
            if not anchor:
                return []
            a_norm = re.sub(r"\s+", " ", anchor.strip()).lower()
            pos = []
            for i in range(start_index, len(doc_paragraphs)):
                para = doc_paragraphs[i]
                if para is None:
                    continue
                para_norm = re.sub(r"\s+", " ", para.strip()).lower()
                if para_norm and re.search(r"(?<!\w)" + re.escape(a_norm) + r"(?!\w)", para_norm, flags=re.IGNORECASE):
                    pos.append(i)
            return pos

        pos_before = find_positions(anchor_before)
        pos_after = find_positions(anchor_after)

        def candidate_ok(val, anchors):
            if not val:
                return False
            v = val.strip()
            if v == "":
                return False
            if self._is_anchor_like(v, anchors):
                return False
            # не считаем нумерацию раздела ("1.", "2.", "3.")
            if re.match(r"^\d+\.", v):
                return False
            # игнорируем служебные заголовки (Заключение, Введение и пр.)
            if any(sw in v.lower() for sw in stop_words):
                return False
            # также игнорируем подписи/инициалы как значения
            if re.search(r"подпись|иниц|инициалы|фамил", v.lower()):
                return False
            if expected_type:
                return self._validate_type(v, expected_type)
            return True

        # Сценарий: есть оба anchors -> ищем между ними (как раньше), но с ограничением дистанции
        if pos_before and pos_after:
            best = None
            best_dist = None
            for b in pos_before:
                for a in pos_after:
                    if b > a:
                        continue
                    dist = a - b
                    if best_dist is None or dist < best_dist:
                        best_dist = dist
                        best = (b, a)
            if best:
                b, a = best
                # ограничение дистанции
                max_doc_distance = 8
                if a - b > max_doc_distance:
                    pos_after = []  # anchor_after слишком далеко
                else:
                    consecutive = (a - b) <= 1
                    if consecutive or next_is_placeholder:
                        return False, None, -1

                    # проверяем текст между anchors
                    for k in range(b + 1, a):
                        mid = doc_paragraphs[k].strip()
                        if candidate_ok(mid, [anchor_before, anchor_after]):
                            return True, mid, k

        # только anchor_before
        if pos_before:
            for b in pos_before:
                para_b = doc_paragraphs[b] or ""
                low_b = para_b.lower()
                ab = anchor_before.strip().lower()
                idx_b = low_b.find(ab)
                if idx_b != -1:
                    after_b = para_b[idx_b + len(ab) :].strip()
                    if candidate_ok(after_b, [anchor_before, anchor_after]):
                        return True, after_b, b
                # ограничиваем поиск максимум до первой пустой строки или до служебного слова
                for k in range(b + 1, len(doc_paragraphs)):
                    cand_raw = doc_paragraphs[k]
                    # если anchor_after отсутствует — пропускаем пустые строки
                    if not anchor_after:
                        if cand_raw is None:
                            continue
                        cand = cand_raw.strip()
                        if cand == "":
                            continue  # просто пропускаем пустые строки
                    else:
                        if cand_raw is None or cand_raw.strip() == "":
                            break

                    cand = cand_raw.strip()
                    # если встретили служебный раздел — прекращаем поиск
                    if any(sw in cand.lower() for sw in stop_words):
                        break
                    # если это явно подпись/инициалы — прекращаем (не берём как значение)
                    if re.search(r"^\(?\s*(подпись|иниц|инициалы|фамил)", cand.lower()):
                        break
                    if candidate_ok(cand, [anchor_before, anchor_after]):
                        return True, cand, k

                    cand = cand_raw.strip()
                    # если встретили служебный раздел — прекращаем поиск
                    if any(sw in cand.lower() for sw in stop_words):
                        break
                    # если это явно подпись/инициалы — прекращаем (не берём как значение)
                    if re.search(r"^\(?\s*(подпись|иниц|инициалы|фамил)", cand.lower()):
                        break
                    if candidate_ok(cand, [anchor_before, anchor_after]):
                        return True, cand, k

        # только anchor_after
        if pos_after:
            for a in pos_after:
                para_a = doc_paragraphs[a] or ""
                low_a = para_a.lower()
                aa = anchor_after.strip().lower()
                idx_a = low_a.find(aa)
                if idx_a != -1:
                    before_a = para_a[:idx_a].strip()
                    if candidate_ok(before_a, [anchor_before, anchor_after]):
                        return True, before_a, a
                # смотрим только одну строчку перед anchor_after (основной сценарий даты/имени)
                k = a - 1
                if k >= 0:
                    cand_raw = doc_paragraphs[k]
                    if cand_raw and cand_raw.strip():
                        cand = cand_raw.strip()
                        # служебные слова / подпись не считаем значением
                        if any(sw in cand.lower() for sw in stop_words):
                            break
                        if re.search(r"^\(?\s*(подпись|иниц|инициалы|фамил)", cand.lower()):
                            break
                        if candidate_ok(cand, [anchor_before, anchor_after]):
                            return True, cand, k

        # ❌ удалён прежний глобальный fallback
        return False, None, -1

    # ------------------------------------------------------------
    def check_document(self, doc_paragraphs: list) -> list:
        results = []
        cursor = 0

        for ph in self.template.get_placeholders():
            name = ph["name"]
            expected_type = ph["type"]
            optional = ph["optional"]
            anchor_before = ph.get("anchor_before", "").strip()
            anchor_after = ph.get("anchor_after", "").strip()
            next_is_placeholder = ph.get("next_is_placeholder", False)

            found, value, found_idx = self._find_value_using_anchors(
                anchor_before, anchor_after, doc_paragraphs, cursor, expected_type, next_is_placeholder
            )

            if not found:
                status = "missing_optional" if optional else "missing"
                results.append(
                    {
                        "field": name,
                        "status": status,
                        "optional": optional,
                        "anchor_before": anchor_before,
                        "anchor_after": anchor_after,
                    }
                )
                continue

            is_valid = self._validate_type(value, expected_type)
            results.append(
                {
                    "field": name,
                    "value": value,
                    "expected_type": expected_type,
                    "status": "ok" if is_valid else "invalid",
                    "optional": optional,
                    "anchor_before": anchor_before,
                    "anchor_after": anchor_after,
                    "found_paragraph_index": found_idx,
                }
            )
            cursor = max(cursor, found_idx + 1)

        return results

    # ------------------------------------------------------------
    def generate_report(self, file_name: str, results: list) -> str:
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        header = (
            f"Отчёт проверки\n"
            f"Файл: {file_name}\n"
            f"Шаблон: {os.path.basename(self.template.source_path)}\n"
            f"Дата: {now}\n\n"
        )

        lines = [header, "=== Проверка заполнителей ===\n"]

        for r in results:
            if r["status"] == "ok":
                lines.append(f"✅ {r['field']} — найдено: {r['value']}")
            elif r["status"] == "invalid":
                lines.append(
                    f"⚠️ {r['field']} — найдено, но тип не соответствует ({r['expected_type']}): {r['value']}"
                )
            elif r["status"] == "missing_optional":
                lines.append(f"ℹ️ {r['field']} — отсутствует (необязательное)")
            elif r["status"] == "missing":
                lines.append(f"❌ {r['field']} — отсутствует или не найдено корректное значение")

        return "\n".join(lines)


# ============================================================
#  GUI
# ============================================================
class AppGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Проверка документа по шаблону")
        self.root.geometry("780x600")

        self.template = None
        self.document_path = None
        self.document_paragraphs = None
        self.checker = None
        self._last_report_text = ""

        self._build_ui()

    def _build_ui(self):
        frame = tk.Frame(self.root, padx=10, pady=10)
        frame.pack(fill="both", expand=True)

        tk.Label(frame, text="Проверка документа по шаблону [[...]]", font=("Segoe UI", 14, "bold")).pack(
            pady=(0, 8)
        )

        btn_frame = tk.Frame(frame)
        btn_frame.pack(fill="x", pady=6)

        self.load_template_btn = tk.Button(btn_frame, text="Загрузить шаблон", width=18, command=self.load_template)
        self.load_template_btn.pack(side="left", padx=(0, 8))

        self.load_doc_btn = tk.Button(btn_frame, text="Загрузить документ", width=22, state="disabled", command=self.load_document)
        self.load_doc_btn.pack(side="left", padx=(0, 8))

        self.run_check_btn = tk.Button(btn_frame, text="Проверить", width=12, state="disabled", command=self.run_check)
        self.run_check_btn.pack(side="left")

        self.save_report_btn = tk.Button(btn_frame, text="Сохранить отчёт", width=16, state="disabled", command=self.save_report)
        self.save_report_btn.pack(side="right")

        self.template_label = tk.Label(frame, text="Шаблон: (не загружен)", anchor="w")
        self.template_label.pack(fill="x")
        self.document_label = tk.Label(frame, text="Документ: (не загружен)", anchor="w")
        self.document_label.pack(fill="x")

        self.result_text = scrolledtext.ScrolledText(frame, wrap=tk.WORD, width=110, height=30, font=("Segoe UI", 10))
        self.result_text.pack(fill="both", expand=True)
        self.result_text.insert(tk.END, "Программа для автоматической проверки самостоятельных работ.\nПожалуйста, загрузите шаблон и документ.")

    def load_template(self):
        path = filedialog.askopenfilename(filetypes=[("Word", "*.docx"), ("PDF", "*.pdf")])
        if not path:
            return
        try:
            tpl = Template.load_from_file(path)
            self.template = tpl
            info = []
            for i, p in enumerate(tpl.get_placeholders(), 1):
                info.append(f"{i}. {p['name']} ({p['type']})")
                info.append(f"   anchor_before: '{p['anchor_before']}'")
                info.append(f"   anchor_after : '{p['anchor_after']}'")
                info.append(f"   para_index: {p['para_index']}\n")
            self.template_label.config(text=f"Шаблон: {os.path.basename(path)} — {len(tpl.get_placeholders())} полей")
            self.result_text.delete(1.0, tk.END)
            self.result_text.insert(tk.END, "\n".join(info))
            self.load_doc_btn.config(state="normal")
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))

    def load_document(self):
        path = filedialog.askopenfilename(filetypes=[("Word", "*.docx"), ("PDF", "*.pdf")])
        if not path:
            return
        try:
            self.document_paragraphs = DocumentLoader.get_paragraphs(path)
            self.document_path = path
            self.document_label.config(text=f"Документ: {os.path.basename(path)}")
            self.run_check_btn.config(state="normal")
        except Exception as e:
            messagebox.showerror("Ошибка загрузки", str(e))

    def run_check(self):
        if not self.template or not self.document_paragraphs:
            messagebox.showwarning("Ошибка", "Сначала загрузите шаблон и документ.")
            return
        self.checker = DocumentChecker(self.template)
        results = self.checker.check_document(self.document_paragraphs)
        report = self.checker.generate_report(os.path.basename(self.document_path), results)
        self.result_text.delete(1.0, tk.END)
        self.result_text.insert(tk.END, report)
        self._last_report_text = report
        self.save_report_btn.config(state="normal")

    def save_report(self):
        if not self._last_report_text:
            return
        path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text", "*.txt")])
        if not path:
            return
        with open(path, "w", encoding="utf-8") as f:
            f.write(self._last_report_text)
        messagebox.showinfo("Готово", f"Отчёт сохранён: {os.path.basename(path)}")


# ============================================================
#  main
# ============================================================
def main():
    root = tk.Tk()
    app = AppGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()
