import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import docx
import PyPDF2
import os
import re
from datetime import datetime

# -------------------------
# DocumentLoader
# -------------------------
class DocumentLoader:
    """Загружает текст и параграфы из .docx и .pdf файлов."""
    @staticmethod
    def _normalize_text(s: str) -> str:
        if s is None:
            return ""
        # заменить NBSP и похожие невидимые символы на обычный пробел
        s = s.replace("\u00A0", " ").replace("\u200B", "").replace("\uFEFF", "")
        # заменить множественные пробелы/переводы строк внутри строки на один пробел
        s = re.sub(r"[ \t\v\f\u00A0]+", " ", s)
        return s.strip()

    @staticmethod
    def load_docx_text_and_paragraphs(path: str):
        doc = docx.Document(path)
        paragraphs = []

        # Обрабатываем обычные параграфы
        for p in doc.paragraphs:
            text = "".join(run.text for run in p.runs)
            text = DocumentLoader._normalize_text(text)
            if text and text.strip():
                paragraphs.append(text)

        # Обрабатываем таблицы, но не добавляем дубли
        seen = set(paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        text = "".join(run.text for run in p.runs)
                        text = DocumentLoader._normalize_text(text)
                        if text and text.strip() and text not in seen:
                            paragraphs.append(text)
                            seen.add(text)

        return "\n".join(paragraphs), paragraphs


    @staticmethod
    def load_pdf_text_and_paragraphs(path: str):
        # PDF извлекается построчно — будем считать строки параграфами
        text_lines = []
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text() or ""
                # splitlines сохраняет структуру
                for ln in page_text.splitlines():
                    ln_norm = DocumentLoader._normalize_text(ln)
                    if ln_norm:
                        text_lines.append(ln_norm)
        full_text = "\n".join(text_lines)
        return full_text, text_lines

    @staticmethod
    def get_text(path: str) -> str:
        lower = path.lower()
        if lower.endswith(".docx"):
            t, _ = DocumentLoader.load_docx_text_and_paragraphs(path)
            return t
        elif lower.endswith(".pdf"):
            t, _ = DocumentLoader.load_pdf_text_and_paragraphs(path)
            return t
        else:
            raise ValueError("Поддерживаются только .docx и .pdf")

    @staticmethod
    def get_paragraphs(path: str):
        lower = path.lower()
        if lower.endswith(".docx"):
            _, paras = DocumentLoader.load_docx_text_and_paragraphs(path)
            return paras
        elif lower.endswith(".pdf"):
            _, paras = DocumentLoader.load_pdf_text_and_paragraphs(path)
            return paras
        else:
            raise ValueError("Поддерживаются только .docx и .pdf")


# -------------------------
# Template
# -------------------------
class Template:
    """
    Хранит список заполнителей (placeholders).
    Для каждого placeholder также сохраняется 'anchor' — текст-метка,
    который используется для поиска значения в заполненном документе.
    """
    def __init__(self, placeholders=None, source_path=None):
        self.placeholders = placeholders or []
        self.source_path = source_path

    def get_placeholders(self):
        """Возвращает список заполнителей"""
        return self.placeholders

    @staticmethod
    def _normalize_type(raw_type: str) -> str:
        t = raw_type.strip().lower()
        if t in ("int", "integer", "number", "num", "float"):
            return "number"
        if t in ("str", "string", "text"):
            return "string"
        if t in ("date", "dt"):
            return "date"
        return t  # оставить как есть, но лучше поддержать только три основных типа

    @staticmethod
    def extract_placeholders_from_paragraphs(paragraphs: list) -> list:
        """
        Улучшенная версия: лучше обрабатывает различные форматы заполнителей
        """
        placeholders = []
        # Более гибкое регулярное выражение для поиска заполнителей
        inline_pattern = re.compile(r"\[\[\s*([^:\]\n]+?)\s*:\s*([^,\]\n]+?)(?:\s*,\s*(optional))?\s*\]\]", flags=re.IGNORECASE)
    
        # Также ищем заполнители в формате с подчеркиванием (как в date_iss в конце документа)
        underline_pattern = re.compile(r"\_\\?\[\[\s*([^:\]\n]+?)\s*:\s*([^,\]\n]+?)(?:\s*,\s*(optional))?\s*\\?\]\]", flags=re.IGNORECASE)
    
        # Паттерны для пропуска
        skip_patterns = [
            "утверждаю", "задание", "введение", "заключение", 
            "список использованных источников", "примерный календарный график",
            "подпись обучающегося", "руководитель курсового проекта"
        ]
    
        for idx, para in enumerate(paragraphs):
            if not para or para.strip() == "":
                continue
            
            # Пропускаем служебные тексты и повторяющиеся элементы
            if any(skip_text in para.lower() for skip_text in skip_patterns):
                continue
            
            # Ищем обычные заполнители
            for m in inline_pattern.finditer(para):
                raw_name = m.group(1).strip()
                raw_type = m.group(2).strip()
                optional_flag = bool(m.group(3))
            
                # Улучшенное определение anchor
                left_part = para[:m.start()].strip()
                if left_part:
                    anchor = left_part
                else:
                    # Ищем предыдущий непустой параграф, но пропускаем служебные
                    anchor = ""
                    for j in range(idx-1, -1, -1):
                        prev_para = paragraphs[j].strip()
                        if prev_para and not any(skip_text in prev_para.lower() for skip_text in skip_patterns):
                            anchor = prev_para
                            break
            
                placeholders.append({
                    "name": raw_name,
                    "type": Template._normalize_type(raw_type),
                    "optional": optional_flag,
                    "anchor": anchor,
                    "source_paragraph": para
                })
        
            # Ищем заполнители в формате с подчеркиванием (как _[[date_iss: string]])
            for m in underline_pattern.finditer(para):
                raw_name = m.group(1).strip()
                raw_type = m.group(2).strip()
                optional_flag = bool(m.group(3))
            
                # Для такого формата anchor обычно текст перед заполнителем в той же строке
                left_part = para[:m.start()].strip()
                if left_part:
                    anchor = left_part
                else:
                    anchor = ""
                    for j in range(idx-1, -1, -1):
                        prev_para = paragraphs[j].strip()
                        if prev_para and not any(skip_text in prev_para.lower() for skip_text in skip_patterns):
                            anchor = prev_para
                            break
            
                placeholders.append({
                    "name": raw_name,
                    "type": Template._normalize_type(raw_type),
                    "optional": optional_flag,
                    "anchor": anchor,
                    "source_paragraph": para
                })
    
        # Улучшенная дедупликация - учитываем имя и точный anchor
        seen = set()
        unique_placeholders = []
    
        for p in placeholders:
            # Нормализуем anchor - убираем лишние пробелы и приводим к нижнему регистру
            normalized_anchor = re.sub(r'\s+', ' ', p["anchor"].strip()).lower() if p["anchor"] else "no_anchor"
            key = (p["name"].lower(), normalized_anchor)
        
            if key not in seen:
                seen.add(key)
                unique_placeholders.append(p)
    
        return unique_placeholders

    @classmethod
    def load_from_file(cls, path: str):
        lower = path.lower()
        if lower.endswith(".docx"):
            _, paragraphs = DocumentLoader.load_docx_text_and_paragraphs(path)
        elif lower.endswith(".pdf"):
            _, paragraphs = DocumentLoader.load_pdf_text_and_paragraphs(path)
        else:
            raise ValueError("Поддерживаются только .docx и .pdf")
        
        # Отладочная информация
        print(f"Извлечено параграфов: {len(paragraphs)}")
        for i, p in enumerate(paragraphs[:10]):  # первые 10 для примера
            if p.strip():
                print(f"Параграф {i}: '{p}'")
        
        placeholders = cls.extract_placeholders_from_paragraphs(paragraphs)
        
        if not placeholders:
            raise ValueError("В шаблоне не найдено ни одного заполнителя [[...]]")
        
        return cls(placeholders=placeholders, source_path=path)


# -------------------------
# DocumentChecker
# -------------------------
class DocumentChecker:
    """Проверяет заполненные значения, опираясь на anchor из шаблона."""
    def __init__(self, template: Template):
        self.template = template

    @staticmethod
    def _validate_type(value: str, expected_type: str) -> bool:
        if value is None:
            return False
        v = value.strip()
        if expected_type == "string":
            return bool(re.search(r"[A-Za-zА-Яа-яЁё]", v))
        elif expected_type == "number":
            return bool(re.fullmatch(r"[+-]?\d+([.,]\d+)?", v))
        elif expected_type == "date":
            # допустимые форматы: 16.09.2025, 16.9.2025, 16 сентября 2025, 16 сент. 2025, 16 сент 2025
            if re.fullmatch(r"\d{1,2}\.\d{1,2}\.\d{4}", v):
                return True
            # русская запись даты: 16 сентября 2025 (слово месяц)
            if re.search(r"\d{1,2}\s+[А-Яа-яёЁ]+\.?\s+\d{4}", v):
                return True
            return False
        else:
            # если неизвестный тип — считаем валидным, но можно поменять на False
            return True

    def _find_value_by_anchor_in_paragraphs(self, anchor: str, doc_paragraphs: list, expected_type: str = None) -> (bool, str):
        """
        Улучшенный поиск значений по anchor.

        Логика:
        1. Если anchor пуст — вернуть (False, None).
        2. Нормализуем anchor (убираем мн. пробелы, приводим к lower).
        3. Ищем в параграфах:
           - сначала — параграфы, где anchor встречается как отдельное слово (word boundary).
             Среди таких вхождений отдаём приоритет тем, где:
               a) после anchor в той же строке есть текст (и это короткий текст / число / дата),
               b) или параграф начинается с anchor (например "Курс:"), тогда смотрим текст после anchor.
           - если таких нет — берём следующий непустой параграф после параграфа с anchor,
             но проверяем на соответствие expected_type (если указан).
        4. Если найдено несколько кандидатов, выбираем наиболее подходящий (например, где value соответствует expected_type,
           или где value короче 40 символов).
        """
        if not anchor:
            return False, None

        # нормализуем anchor для поиска
        anchor_norm = re.sub(r'\s+', ' ', anchor.strip())
        anchor_escaped = re.escape(anchor_norm)
        # собираем паттерн для "как отдельное слово" — добавляем границы \b, но для Unicode нам пригодится
        # использовать (?<!\w) и (?!\w) — чтобы работать с кириллицей корректно
        word_pat = re.compile(r"(?<!\w)" + anchor_escaped + r"(?!\w)", flags=re.IGNORECASE)

        candidates = []  # tuples (priority, para_index, value_candidate, reason)

        for i, para in enumerate(doc_paragraphs):
            if not para:
                continue
            para_norm = re.sub(r'\s+', ' ', para.strip())

            # 1) ищем anchor как отдельное слово
            for m in word_pat.finditer(para_norm):
                # position в нормализованном параграфе
                start_pos = m.start()
                end_pos = m.end()

                # получим часть параграфа после anchor (оригинальную, не нормализованную)
                # найдём соответствующую позицию в оригинальном параграфе
                para_original = doc_paragraphs[i]
                # простая эвристика: найдём первое вхождение anchor_norm (case-insensitive) в оригинале
                idx_orig = para_original.lower().find(anchor_norm.lower())
                if idx_orig != -1:
                    after = para_original[idx_orig + len(anchor_norm):].strip()
                else:
                    after = para_original[end_pos:].strip()  # fallback

                # Оцениваем кандидатуру:
                # Если есть текст сразу после anchor в том же параграфе — это хороший кандидат.
                if after:
                    # приоритет выше, если:
                    # - параграф начинается с anchor (anchor в начале) OR
                    # - текст после anchor короткий (<=40) OR соответствует expected_type
                    starts_with = para_norm.lower().startswith(anchor_norm.lower())
                    length_score = 0 if len(after) <= 40 else 1
                    type_matches = False
                    if expected_type:
                        type_matches = self._validate_type(after, expected_type)
                    priority = 10  # базовый
                    if starts_with:
                        priority -= 4
                    if type_matches:
                        priority -= 3
                    if length_score == 1:
                        priority += 1
                    candidates.append((priority, i, after, f"same_para (after)"))
                else:
                    # если после anchor нет текста в той же строке — возможный кандидат: следующий непустой параграф
                    # оценим его ниже
                    # пометим как потенциальный и обработаем дальше
                    # priority выше, если expected_type matches next paragraph
                    next_val = None
                    for j in range(i+1, len(doc_paragraphs)):
                        nxt = doc_paragraphs[j].strip()
                        if not nxt:
                            continue
                        # не брать следующий параграф, если он повторяет сам anchor (зацикливание)
                        if anchor_norm.lower() == re.sub(r'\s+', ' ', nxt.lower()):
                            continue
                        next_val = nxt
                        break
                    if next_val:
                        type_matches = False
                        if expected_type:
                            type_matches = self._validate_type(next_val, expected_type)
                        priority = 14
                        if type_matches:
                            priority -= 5
                        # если следующий параграф короткий — лучше
                        if len(next_val) <= 40:
                            priority -= 1
                        candidates.append((priority, i+1, next_val, f"next_para"))

        # Если нашли кандидатов — выбираем минимальный priority (лучший)
        if candidates:
            candidates.sort(key=lambda x: (x[0], x[1]))  # сначала priority, затем порядковый номер
            best = candidates[0]
            return True, best[2]

        # Если не нашли, попытаться ещё раз более ленно: искать anchor как substring (fallback),
        # но только если anchor длинный (>3) — чтобы не ловить короткие слова везде.
        if len(anchor_norm) > 3:
            for i, para in enumerate(doc_paragraphs):
                if not para:
                    continue
                para_norm = re.sub(r'\s+', ' ', para.strip())
                if anchor_norm.lower() in para_norm.lower():
                    # попробуем взять текст после первого вхождения
                    idx = para_norm.lower().find(anchor_norm.lower())
                    para_original = doc_paragraphs[i]
                    idx_orig = para_original.lower().find(anchor_norm.lower())
                    after = para_original[idx_orig + len(anchor_norm):].strip() if idx_orig != -1 else ""
                    if after:
                        if expected_type is None or self._validate_type(after, expected_type):
                            return True, after
                    # else попробовать следующий параграф
                    for j in range(i+1, len(doc_paragraphs)):
                        nxt = doc_paragraphs[j].strip()
                        if nxt:
                            if expected_type is None or self._validate_type(nxt, expected_type):
                                return True, nxt
                            else:
                                break

        return False, None


    def check_document(self, doc_paragraphs: list) -> list:
        """
        Для каждого placeholder:
         - используем anchor (из шаблона) и doc_paragraphs (список параграфов из проверяемого документа),
         - извлекаем значение и проверяем тип.
        """
        results = []
        for ph in self.template.get_placeholders():
            name = ph["name"]
            expected_type = ph["type"]
            optional = ph["optional"]
            anchor = ph.get("anchor", "").strip()

            found, value = self._find_value_by_anchor_in_paragraphs(anchor, doc_paragraphs, expected_type=expected_type)
            if not found:
                if optional:
                    results.append({"field": name, "status": "missing_optional", "optional": True})
                else:
                    results.append({"field": name, "status": "missing", "optional": False})
                continue

            is_valid = self._validate_type(value, expected_type)
            results.append({
                "field": name,
                "value": value,
                "expected_type": expected_type,
                "status": "ok" if is_valid else "invalid",
                "optional": optional,
                "anchor": anchor
            })
        return results

    def generate_report(self, file_name: str, results: list) -> str:
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        header = (
            f"Отчёт проверки\n"
            f"Файл: {file_name}\n"
            f"Шаблон: {os.path.basename(self.template.source_path) if self.template.source_path else '(не указан)'}\n"
            f"Дата: {now}\n\n"
        )
    
        # Группируем результаты по имени поля
        grouped_results = {}
        for r in results:
            field_name = r["field"]
            if field_name not in grouped_results:
                grouped_results[field_name] = []
            grouped_results[field_name].append(r)
    
        lines = [header, "=== Проверка заполнителей ===\n"]
    
        for field_name, field_results in grouped_results.items():
            if len(field_results) == 1:
                # Один результат для поля
                r = field_results[0]
                if r["status"] == "ok":
                    lines.append(f"✅ {r['field']} — корректно ({r['expected_type']}): {r.get('value','')}")
                elif r["status"] == "invalid":
                    lines.append(f"⚠️ {r['field']} — найдено, но тип не соответствует ({r['expected_type']}): {r.get('value','')}")
                elif r["status"] == "missing_optional":
                    lines.append(f"ℹ️ {r['field']} — отсутствует (необязательное)")
                elif r["status"] == "missing":
                    lines.append(f"❌ {r['field']} — отсутствует (обязательное)")
            else:
                # Несколько результатов для одного поля
                lines.append(f"🔍 {field_name} — найдено {len(field_results)} вхождений:")
                for i, r in enumerate(field_results, 1):
                    anchor_info = f" (anchor: «{r.get('anchor', '')}»)" if r.get('anchor') else ""
                    if r["status"] == "ok":
                        lines.append(f"   {i}. ✅ Корректно: {r.get('value','')}{anchor_info}")
                    elif r["status"] == "invalid":
                        lines.append(f"   {i}. ⚠️ Тип не соответствует: {r.get('value','')}{anchor_info}")
                    elif r["status"] == "missing_optional":
                        lines.append(f"   {i}. ℹ️ Отсутствует (необязательное){anchor_info}")
                    elif r["status"] == "missing":
                        lines.append(f"   {i}. ❌ Отсутствует (обязательное){anchor_info}")
                lines.append("")  # пустая строка для разделения
    
        return "\n".join(lines)


# -------------------------
# AppGUI (Tkinter)
# -------------------------
class AppGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Проверка документа по шаблону [[...]] (anchor-based)")
        self.root.geometry("760x560")
        self.template = None
        self.document_path = None
        self.document_paragraphs = None
        self.checker = None
        self._build_ui()

    def _build_ui(self):
        frame = tk.Frame(self.root, padx=10, pady=10)
        frame.pack(fill="both", expand=True)

        title = tk.Label(frame, text="Проверка документа по шаблону [[...]]", font=("Segoe UI", 14, "bold"))
        title.pack(pady=(0,8))

        btn_frame = tk.Frame(frame)
        btn_frame.pack(fill="x", pady=6)

        self.load_template_btn = tk.Button(btn_frame, text="Загрузить шаблон", command=self.load_template, width=18)
        self.load_template_btn.pack(side="left", padx=(0,8))

        self.load_doc_btn = tk.Button(btn_frame, text="Загрузить документ", command=self.load_document, width=22, state="disabled")
        self.load_doc_btn.pack(side="left", padx=(0,8))

        self.run_check_btn = tk.Button(btn_frame, text="Проверить", command=self.run_check, width=12, state="disabled")
        self.run_check_btn.pack(side="left")

        self.save_report_btn = tk.Button(btn_frame, text="Сохранить отчёт", command=self.save_report, width=16, state="disabled")
        self.save_report_btn.pack(side="right")

        info_frame = tk.Frame(frame)
        info_frame.pack(fill="x", pady=(6, 8))
        self.template_label = tk.Label(info_frame, text="Шаблон: (не загружен)", anchor="w")
        self.template_label.pack(fill="x")
        self.document_label = tk.Label(info_frame, text="Документ: (не загружен)", anchor="w")
        self.document_label.pack(fill="x")

        self.result_text = scrolledtext.ScrolledText(frame, wrap=tk.WORD, width=100, height=26, font=("Segoe UI", 10))
        self.result_text.pack(fill="both", expand=True)

        welcome = (
            "Инструкция:\n"
            "1) В шаблоне используйте заполнители в формате [[имя_поля: тип]] или [[имя_поля: тип, optional]].\n"
            "   Допустимые (рекомендованные) типы: string, number, date\n"
            "   (поддерживаются также: int/integer/num/float → number; str → string).\n"
            "2) Если поле на отдельной строке — напишите перед ним метку (anchor) на предыдущей строке, например:\n"
            "   Обучающемуся\n"
            "   [[full_name: string]]\n"
            "   В этом случае anchor = 'Обучающемуся' и система возьмёт следующее непустое значение в заполненном документе.\n"
            "3) Загрузите шаблон, затем документ и нажмите «Проверить».\n"
        )
        self.result_text.insert(tk.END, welcome)
        self._last_report_text = ""

    def load_template(self):
        path = filedialog.askopenfilename(title="Выберите шаблон", filetypes=[("Документы Word", "*.docx"), ("PDF файлы", "*.pdf")])
        if not path:
            return
        try:
            tpl = Template.load_from_file(path)
            self.template = tpl
        
            # Детальная информация о найденных заполнителях
            info_lines = []
            for i, p in enumerate(tpl.get_placeholders()):
                info_lines.append(f"{i+1}. {p['name']} ({p['type']}{', optional' if p['optional'] else ''})")
                info_lines.append(f"   Anchor: '{p['anchor']}'")
                info_lines.append(f"   Источник: '{p.get('source_paragraph', '')[:50]}...'")
                info_lines.append("")
        
            self.template_label.config(text=f"Шаблон: {os.path.basename(path)} — {len(tpl.get_placeholders())} полей")
            self.result_text.delete(1.0, tk.END)
            self.result_text.insert(tk.END, f"Шаблон загружен: {os.path.basename(path)}\n\nНайденные поля:\n" + "\n".join(info_lines))
        
            self.load_doc_btn.config(state="normal")
            self.run_check_btn.config(state="disabled")
            self.save_report_btn.config(state="disabled")
            self._last_report_text = ""
        
        except Exception as e:
            messagebox.showerror("Ошибка загрузки шаблона", f"{str(e)}\n\nУбедитесь, что:\n1. Файл не поврежден\n2. Заполнители в формате [[name: type]]\n3. Файл не защищен паролем")

    def load_document(self):
        path = filedialog.askopenfilename(title="Выберите документ", filetypes=[("Документы Word", "*.docx"), ("PDF файлы", "*.pdf")])
        if not path:
            return
        try:
            paras = DocumentLoader.get_paragraphs(path)
            # normalize paragraphs (already done inside loader)
            self.document_paragraphs = [p for p in paras]
            self.document_path = path
            self.document_label.config(text=f"Документ: {os.path.basename(path)}")
            self.result_text.delete(1.0, tk.END)
            self.result_text.insert(tk.END, f"Документ загружен: {os.path.basename(path)}\n\nНажмите «Проверить».")
            if self.template:
                self.checker = DocumentChecker(self.template)
                self.run_check_btn.config(state="normal")
        except Exception as e:
            messagebox.showerror("Ошибка загрузки документа", str(e))

    def run_check(self):
        if not self.template or not self.document_paragraphs:
            messagebox.showwarning("Ошибка", "Сначала загрузите шаблон и документ.")
            return
        try:
            results = self.checker.check_document(self.document_paragraphs)
            report = self.checker.generate_report(os.path.basename(self.document_path), results)
            self._last_report_text = report
            self.result_text.delete(1.0, tk.END)
            self.result_text.insert(tk.END, report)
            self.save_report_btn.config(state="normal")
        except Exception as e:
            messagebox.showerror("Ошибка при проверке", str(e))

    def save_report(self):
        if not self._last_report_text:
            messagebox.showinfo("Нет отчёта", "Сначала выполните проверку.")
            return
        path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Текстовый файл", "*.txt")])
        if not path:
            return
        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(self._last_report_text)
            messagebox.showinfo("Готово", f"Отчёт сохранён: {os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("Ошибка сохранения", str(e))


def main():
    root = tk.Tk()
    app = AppGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()
